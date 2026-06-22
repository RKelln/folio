"""Tests for the corpus metadata stripper.

``strip_metadata`` removes authoring metadata (a PII risk) from rendered binary
artifacts; ``read_metadata`` reads it back for verification. Tests for formats
whose required tool is absent are skipped rather than failed.
"""

from __future__ import annotations

import shutil

import pytest

from folio.core.corpus.metadata import read_metadata, strip_metadata

_HAS_EXIFTOOL = shutil.which("exiftool") is not None


def _make_docx(path, author="Confidential Author"):
    import docx

    doc = docx.Document()
    doc.add_heading("Title", level=1)
    doc.add_paragraph("Body text.")
    cp = doc.core_properties
    cp.author = author
    cp.last_modified_by = author
    cp.title = "Secret Title"
    doc.save(str(path))


def _make_xlsx(path, author="Confidential Author"):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws["A1"] = "value"
    wb.properties.creator = author
    wb.properties.lastModifiedBy = author
    wb.properties.title = "Secret Title"
    wb.save(str(path))


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


def _xlsx_available() -> bool:
    try:
        import openpyxl  # noqa: F401

        return True
    except ImportError:
        return False


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
class TestStripDocx:
    def test_strip_clears_author(self, tmp_path):
        if not _docx_available():
            pytest.skip("python-docx not available")
        path = tmp_path / "doc.docx"
        _make_docx(path)
        assert strip_metadata(path) is True
        meta = read_metadata(path)
        assert not meta.get("author")
        assert not meta.get("last_modified_by")

    def test_read_metadata_returns_dict(self, tmp_path):
        if not _docx_available():
            pytest.skip("python-docx not available")
        path = tmp_path / "doc.docx"
        _make_docx(path, author="Jane Doe")
        meta = read_metadata(path)
        assert isinstance(meta, dict)
        assert meta.get("author") == "Jane Doe"


# --------------------------------------------------------------------------- #
# XLSX
# --------------------------------------------------------------------------- #
class TestStripXlsx:
    def test_strip_clears_creator(self, tmp_path):
        if not _xlsx_available():
            pytest.skip("openpyxl not available")
        path = tmp_path / "book.xlsx"
        _make_xlsx(path)
        assert strip_metadata(path) is True
        meta = read_metadata(path)
        assert not meta.get("creator")
        assert not meta.get("last_modified_by")

    def test_read_metadata_returns_creator(self, tmp_path):
        if not _xlsx_available():
            pytest.skip("openpyxl not available")
        path = tmp_path / "book.xlsx"
        _make_xlsx(path, author="John Roe")
        meta = read_metadata(path)
        assert meta.get("creator") == "John Roe"


# --------------------------------------------------------------------------- #
# PDF (needs exiftool + a rendered pdf)
# --------------------------------------------------------------------------- #
@pytest.mark.integration
class TestStripPdf:
    def test_strip_removes_producer_and_author(self, tmp_path):
        from folio.adapters.renderers import PdfRenderer

        r = PdfRenderer()
        if not r.available():
            pytest.skip("pandoc/typst not available to render a PDF")
        if not _HAS_EXIFTOOL:
            pytest.skip("exiftool not available")
        pdf = r.render("# Title\n\nBody paragraph.\n", {}, tmp_path / "doc.pdf")
        assert strip_metadata(pdf) is True
        meta = read_metadata(pdf)
        for leaky in ("Author", "Producer", "Creator", "CreatorTool"):
            assert not meta.get(leaky), f"{leaky} leaked: {meta.get(leaky)!r}"


# --------------------------------------------------------------------------- #
# Missing-tool behaviour
# --------------------------------------------------------------------------- #
class TestMissingExiftool:
    def test_pdf_strip_raises_when_exiftool_absent(self, tmp_path, monkeypatch):
        import folio.core.corpus.metadata as md

        monkeypatch.setattr(md.shutil, "which", lambda _name: None)
        pdf = tmp_path / "fake.pdf"
        pdf.write_bytes(b"%PDF-1.4 fake")
        with pytest.raises(RuntimeError, match="exiftool"):
            strip_metadata(pdf)

    def test_unknown_suffix_returns_false(self, tmp_path):
        odd = tmp_path / "thing.bin"
        odd.write_bytes(b"\x00\x01\x02")
        assert strip_metadata(odd) is False
