"""Tests for the corpus format renderers (Markdown golden → DOCX/XLSX/PDF).

Every renderer degrades gracefully: if its required tools/deps are absent the
relevant test is skipped (never failed). In a fully-provisioned environment
(pandoc + typst + python-docx + openpyxl + poppler + Pillow) every test runs.
"""

from __future__ import annotations

import pytest

from folio.adapters.renderers import (
    DocxRenderer,
    PdfRenderer,
    ScannedPdfRenderer,
    XlsxRenderer,
    available_formats,
    get_renderer,
)
from folio.adapters.renderers.base import Renderer, parse_markdown

SAMPLE_MD = """\
# Operating Grant Application

## Organization Profile

We are a synthetic artist-run centre serving an imaginary community.

### Revenue

| Source | Amount |
| --- | --- |
| **Earned revenue** | $12,500 |
| Grants | $40,000 |

### Expenses

| Category | Amount |
| --- | --- |
| Programming | $30,000 |
| Administration | $22,500 |
"""


# --------------------------------------------------------------------------- #
# Factory
# --------------------------------------------------------------------------- #
class TestRendererFactory:
    def test_docx(self):
        assert isinstance(get_renderer("docx"), DocxRenderer)

    def test_xlsx(self):
        assert isinstance(get_renderer("xlsx"), XlsxRenderer)

    def test_pdf(self):
        assert isinstance(get_renderer("pdf"), PdfRenderer)

    def test_pdf_scanned(self):
        assert isinstance(get_renderer("pdf_scanned"), ScannedPdfRenderer)

    @pytest.mark.parametrize("bogus", ["png", "html", "", "DOCX"])
    def test_unknown_format_raises_valueerror(self, bogus):
        with pytest.raises(ValueError, match="Unknown output format"):
            get_renderer(bogus)

    def test_each_renderer_reports_its_format(self):
        assert get_renderer("docx").output_format == "docx"
        assert get_renderer("xlsx").output_format == "xlsx"
        assert get_renderer("pdf").output_format == "pdf"
        assert get_renderer("pdf_scanned").output_format == "pdf_scanned"

    def test_available_formats_is_subset_of_known(self):
        known = {"docx", "xlsx", "pdf", "pdf_scanned"}
        assert set(available_formats()).issubset(known)


# --------------------------------------------------------------------------- #
# Shared markdown parser
# --------------------------------------------------------------------------- #
class TestParseMarkdown:
    def test_parses_headings_paragraphs_and_tables(self):
        blocks = parse_markdown(SAMPLE_MD)
        kinds = [b.__class__.__name__ for b in blocks]
        assert "MdHeading" in kinds
        assert "MdParagraph" in kinds
        assert "MdTable" in kinds

    def test_table_strips_bold_and_keeps_currency(self):
        blocks = parse_markdown(SAMPLE_MD)
        tables = [b for b in blocks if b.__class__.__name__ == "MdTable"]
        assert len(tables) == 2
        revenue = tables[0]
        assert revenue.section == "Revenue"
        # header + 2 data rows
        assert revenue.rows[0] == ["Source", "Amount"]
        assert ["Earned revenue", "$12,500"] in revenue.rows
        # no '**' survived in any cell
        flat = [c for row in revenue.rows for c in row]
        assert all("**" not in c for c in flat)

    def test_separator_row_is_not_emitted_as_data(self):
        blocks = parse_markdown(SAMPLE_MD)
        tables = [b for b in blocks if b.__class__.__name__ == "MdTable"]
        for t in tables:
            assert not any(set(c) <= {"-", ":", " "} and c for row in t.rows for c in row)

    def test_empty_markdown_yields_no_blocks(self):
        assert parse_markdown("") == []


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
class TestDocxRenderer:
    def test_is_renderer(self):
        assert isinstance(DocxRenderer(), Renderer)

    def test_render(self, tmp_path):
        r = DocxRenderer()
        if not r.available():
            pytest.skip("python-docx not available")
        out = r.render(SAMPLE_MD, {}, tmp_path / "app.docx")
        assert out.exists() and out.stat().st_size > 0

        import docx

        doc = docx.Document(str(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Operating Grant Application" in all_text
        assert "synthetic artist-run centre" in all_text
        assert len(doc.tables) == 2
        # bold markers stripped inside table cells
        cells = [c.text for t in doc.tables for row in t.rows for c in row.cells]
        assert "Earned revenue" in cells
        assert all("**" not in c for c in cells)

    def test_render_clears_author(self, tmp_path):
        r = DocxRenderer()
        if not r.available():
            pytest.skip("python-docx not available")
        out = r.render(SAMPLE_MD, {"author": "Real Person"}, tmp_path / "x.docx")
        import docx

        cp = docx.Document(str(out)).core_properties
        assert not cp.author
        assert not cp.last_modified_by


# --------------------------------------------------------------------------- #
# XLSX
# --------------------------------------------------------------------------- #
class TestXlsxRenderer:
    def test_is_renderer(self):
        assert isinstance(XlsxRenderer(), Renderer)

    def test_render(self, tmp_path):
        r = XlsxRenderer()
        if not r.available():
            pytest.skip("openpyxl not available")
        out = r.render(SAMPLE_MD, {}, tmp_path / "budget.xlsx")
        assert out.exists() and out.stat().st_size > 0

        import openpyxl

        wb = openpyxl.load_workbook(str(out))
        assert "Revenue" in wb.sheetnames
        assert "Expenses" in wb.sheetnames
        rev = wb["Revenue"]
        assert rev["A1"].value == "Source"
        assert rev["B1"].value == "Amount"
        # currency kept as text, bold stripped
        col_a = [c.value for c in rev["A"]]
        assert "Earned revenue" in col_a
        assert any(v == "$12,500" for row in rev.iter_rows(values_only=True) for v in row)

    def test_render_clears_creator(self, tmp_path):
        r = XlsxRenderer()
        if not r.available():
            pytest.skip("openpyxl not available")
        out = r.render(SAMPLE_MD, {"author": "Real Person"}, tmp_path / "y.xlsx")
        import openpyxl

        props = openpyxl.load_workbook(str(out)).properties
        assert not props.creator
        assert not props.lastModifiedBy


# --------------------------------------------------------------------------- #
# PDF (integration — needs pandoc + typst)
# --------------------------------------------------------------------------- #
@pytest.mark.integration
class TestPdfRenderer:
    def test_is_renderer(self):
        assert isinstance(PdfRenderer(), Renderer)

    def test_render(self, tmp_path):
        r = PdfRenderer()
        if not r.available():
            pytest.skip("pandoc and/or typst not available")
        out = r.render(SAMPLE_MD, {}, tmp_path / "app.pdf")
        assert out.exists()
        data = out.read_bytes()
        assert data[:4] == b"%PDF"
        assert len(data) > 4


@pytest.mark.integration
class TestScannedPdfRenderer:
    def test_is_renderer(self):
        assert isinstance(ScannedPdfRenderer(), Renderer)

    def test_render_image_only_pdf(self, tmp_path):
        r = ScannedPdfRenderer()
        if not r.available():
            pytest.skip("pandoc/typst/poppler/Pillow not all available")
        out = r.render(SAMPLE_MD, {}, tmp_path / "scanned.pdf")
        assert out.exists()
        data = out.read_bytes()
        assert data[:4] == b"%PDF"
        # Image-only render must not contain extractable selectable text.
        import shutil
        import subprocess

        pdftotext = shutil.which("pdftotext")
        if pdftotext:
            extracted = subprocess.run(
                [pdftotext, "-q", str(out), "-"],
                capture_output=True,
                check=False,
            ).stdout.decode("utf-8", "replace")
            assert "Operating Grant Application" not in extracted
