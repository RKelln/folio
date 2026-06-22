"""DOCX renderer — golden Markdown → Word document via python-docx.

Headings (``#``/``##``/``###`` ...) become Word heading styles, paragraphs
become body paragraphs, and GitHub pipe tables become Word tables. Document
core properties are cleared before save (using the single canonical scrub list
``folio.core.corpus.metadata.DOCX_CORE_PROPS``) so no authoring metadata —
author, last-modified-by, timestamps, revision, etc. — leaks into the corpus.
"""

from __future__ import annotations

import logging
from pathlib import Path

from folio.adapters.renderers.base import (
    MdHeading,
    MdParagraph,
    MdTable,
    Renderer,
    parse_markdown,
)
from folio.core.corpus.metadata import DOCX_CORE_PROPS

logger = logging.getLogger(__name__)


class DocxRenderer(Renderer):
    """Render Markdown to a ``.docx`` file using python-docx."""

    name = "python-docx"
    output_format = "docx"

    def available(self) -> bool:
        try:
            import docx  # type: ignore[import-not-found]  # noqa: F401

            return True
        except ImportError as exc:
            logger.warning("DocxRenderer unavailable: python-docx not importable (%s)", exc)
            return False

    def render(self, markdown: str, meta: dict, out_path: Path) -> Path:
        try:
            import docx  # type: ignore[import-not-found]
        except ImportError as exc:
            raise RuntimeError(
                "python-docx not installed; cannot render .docx. "
                "Install with: pip install python-docx"
            ) from exc

        out_path = Path(out_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

        document = docx.Document()
        for block in parse_markdown(markdown):
            if isinstance(block, MdHeading):
                # python-docx supports heading levels 0-9; clamp ours to 1-6.
                document.add_heading(block.text, level=min(block.level, 6))
            elif isinstance(block, MdParagraph):
                document.add_paragraph(block.text)
            elif isinstance(block, MdTable):
                self._add_table(document, block)

        self._clear_core_properties(document)
        document.save(str(out_path))
        return out_path

    @staticmethod
    def _add_table(document, table: MdTable) -> None:
        """Append a Word table built from a parsed MdTable (header + rows)."""
        if not table.rows:
            return
        ncols = max(len(r) for r in table.rows)
        word_table = document.add_table(rows=0, cols=ncols)
        word_table.style = "Table Grid"
        for row in table.rows:
            cells = word_table.add_row().cells
            for idx in range(ncols):
                cells[idx].text = row[idx] if idx < len(row) else ""

    @staticmethod
    def _clear_core_properties(document) -> None:
        """Blank out every authoring core property before saving.

        Applies the single canonical scrub list
        :data:`folio.core.corpus.metadata.DOCX_CORE_PROPS` so this renderer and
        the post-render ``strip_metadata`` stripper stay in lock-step. Each set
        is guarded: an unsettable property is logged, never fatal.
        """
        cp = document.core_properties
        for attr, value in DOCX_CORE_PROPS:
            try:
                setattr(cp, attr, value)
            except (ValueError, TypeError, AttributeError) as exc:
                logger.warning("Could not clear docx core property %s: %s", attr, exc)
