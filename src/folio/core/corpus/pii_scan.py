"""PII scan gate for the synthetic grant corpus.

PURPOSE
    Provide the safety mechanism that MUST pass before any corpus file is
    committed. The synthetic corpus mirrors the *structure* of real grant
    archives but its content must contain no personally identifiable
    information (PII). This module scans text and files for structural PII
    (emails, phone numbers, Canadian SINs, US SSNs, postal/ZIP codes, currency
    amounts) and for denylisted real names (loaded from config, per AGENTS.md).

DESIGN
    Detection is deliberately conservative: it prefers false positives over
    misses, because a missed leak is far worse than a noisy gate. Regexes are
    anchored with digit/word boundaries so common non-PII (years, dates, plain
    counts) do not trip the numeric detectors.

PUBLIC API
    load_denylist(path)        -> list[str]
    Finding (dataclass)        -> kind, match, line, context
    PIIReport (dataclass)      -> path, findings, .clean
    scan_text(text, denylist)  -> list[Finding]
    scan_file(path, denylist)  -> PIIReport
    scan_paths(paths, denylist)-> list[PIIReport]

CONFIG-DRIVEN
    The name denylist is configuration (templates/corpus/pii-denylist.yaml or a
    per-org override), never hardcoded here.
"""

from __future__ import annotations

import logging
import re
import shutil
import subprocess
from collections.abc import Iterable
from dataclasses import dataclass, field
from importlib import resources
from pathlib import Path

import yaml

logger = logging.getLogger(__name__)

# Text formats read directly as UTF-8.
_TEXT_SUFFIXES = {".md", ".markdown", ".txt", ".html", ".htm", ".csv"}

# Width (characters) of context shown on either side of a match.
_CONTEXT_WIDTH = 40


# --------------------------------------------------------------------------- #
# Detection patterns
# --------------------------------------------------------------------------- #
# Each numeric pattern uses (?<!\d) / (?!\d) guards so it cannot match a
# fragment of a longer run of digits (which keeps phone/SIN/SSN/ZIP distinct).
_EMAIL = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")

_PHONE = re.compile(
    r"(?<!\d)(?:\+?1[ .\-]?)?\(?\d{3}\)?[ .\-]?\d{3}[ .\-]?\d{4}(?!\d)"
)

# Canadian SIN: 9 digits, grouped 3-3-3 with optional space/dash separators.
_SIN = re.compile(r"(?<!\d)\d{3}[ \-]?\d{3}[ \-]?\d{3}(?!\d)")

# US SSN: strict 3-2-4 dashed grouping.
_SSN = re.compile(r"(?<!\d)\d{3}-\d{2}-\d{4}(?!\d)")

# Canadian postal code: A1A 1A1, optional space or dash, also no-separator form.
_POSTAL = re.compile(r"\b[A-Za-z]\d[A-Za-z][ \-]?\d[A-Za-z]\d\b")

# US ZIP+4: five digits, dash, four digits.
_ZIP4 = re.compile(r"(?<!\d)\d{5}-\d{4}(?!\d)")

# Currency: $ / CAD / USD prefix followed by a (optionally grouped) amount.
_CURRENCY = re.compile(
    r"(?:\$|\bCAD\b|\bUSD\b)[ ]?\d+(?:,\d{3})*(?:\.\d{1,2})?"
)

# (kind, compiled pattern) in scan order. Postal must precede nothing special;
# order only affects the sequence of findings, not correctness.
_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("email", _EMAIL),
    ("phone", _PHONE),
    ("ssn", _SSN),
    ("sin", _SIN),
    ("postal_code", _POSTAL),
    ("postal_code", _ZIP4),
    ("currency", _CURRENCY),
]


@dataclass
class Finding:
    """A single PII detection.

    Attributes:
        kind: One of "email", "phone", "sin", "ssn", "postal_code",
            "currency", "denylisted_name", or "unscannable" (the latter is
            recorded when a file's text could not be extracted for scanning).
        match: The exact substring (or extractor error reason) that triggered
            the finding.
        line: 1-based line number within the scanned text.
        context: A short snippet of the surrounding line, for human review.
    """

    kind: str
    match: str
    line: int
    context: str


@dataclass
class PIIReport:
    """Result of scanning a single file.

    Attributes:
        path: The scanned file path (as a string).
        findings: All PII findings; an empty list means the file is clean.
    """

    path: str
    findings: list[Finding] = field(default_factory=list)

    @property
    def clean(self) -> bool:
        """True when no findings were recorded (the file may be committed)."""
        return not self.findings


# --------------------------------------------------------------------------- #
# Denylist loading
# --------------------------------------------------------------------------- #
def _bundled_denylist_path() -> Path:
    """Resolve the path to the bundled default denylist YAML.

    Tries importlib.resources first (works for installed packages) and falls
    back to a path relative to this module (works in an editable/source tree).
    """
    try:
        traversable = resources.files("folio.templates").joinpath(
            "corpus", "pii-denylist.yaml"
        )
        candidate = Path(str(traversable))
        if candidate.exists():
            return candidate
    except (ModuleNotFoundError, FileNotFoundError) as exc:
        logger.debug("importlib.resources lookup failed, falling back: %s", exc)
    # parents[2] == folio package root (corpus -> core -> folio).
    return Path(__file__).resolve().parents[2] / "templates" / "corpus" / "pii-denylist.yaml"


def load_denylist(path: str | Path | None = None) -> list[str]:
    """Load the name denylist from a YAML file with a top-level ``names:`` list.

    Inputs:
        path: Explicit YAML path. If None, the bundled default shipped at
            ``folio/templates/corpus/pii-denylist.yaml`` is loaded.

    Output:
        A list of name strings. Returns an empty list if the file has no
        ``names`` key (but raises if the file itself is missing/invalid).

    Side effects:
        Reads from disk.
    """
    resolved = Path(path) if path is not None else _bundled_denylist_path()
    if not resolved.exists():
        raise FileNotFoundError(
            f"PII denylist not found: {resolved}. The folio package may be "
            f"corrupted, or an explicit path was wrong."
        )
    try:
        data = yaml.safe_load(resolved.read_text(encoding="utf-8")) or {}
    except yaml.YAMLError as exc:
        raise ValueError(f"PII denylist {resolved} is not valid YAML: {exc}") from exc
    names = data.get("names", []) if isinstance(data, dict) else []
    if not isinstance(names, list):
        raise ValueError(
            f"PII denylist {resolved}: 'names' must be a list, got {type(names).__name__}"
        )
    return [str(n) for n in names if str(n).strip()]


def _compile_denylist(names: list[str]) -> list[tuple[str, re.Pattern[str]]]:
    """Compile each denylisted name into a case-insensitive whole-word pattern."""
    compiled: list[tuple[str, re.Pattern[str]]] = []
    for name in names:
        cleaned = name.strip()
        if not cleaned:
            continue
        compiled.append(
            (cleaned, re.compile(rf"\b{re.escape(cleaned)}\b", re.IGNORECASE))
        )
    return compiled


# --------------------------------------------------------------------------- #
# Scanning
# --------------------------------------------------------------------------- #
def _snippet(line: str, start: int, end: int) -> str:
    """Build a short, human-readable context snippet around a match span."""
    lo = max(0, start - _CONTEXT_WIDTH)
    hi = min(len(line), end + _CONTEXT_WIDTH)
    prefix = "…" if lo > 0 else ""
    suffix = "…" if hi < len(line) else ""
    return (prefix + line[lo:hi] + suffix).strip()


def scan_text(text: str, denylist: list[str] | None = None) -> list[Finding]:
    """Scan raw text for PII and return all findings.

    Inputs:
        text: The text to scan.
        denylist: Names to flag (case-insensitive, whole-word). If None, the
            bundled default denylist is loaded so name scanning is always on.

    Output:
        A list of Finding objects, ordered by line then detector. Empty list
        means the text is clean.
    """
    if denylist is None:
        denylist = load_denylist()
    name_patterns = _compile_denylist(denylist)

    findings: list[Finding] = []
    for lineno, line in enumerate(text.splitlines(), start=1):
        for kind, pattern in _PATTERNS:
            for m in pattern.finditer(line):
                findings.append(
                    Finding(
                        kind=kind,
                        match=m.group(0),
                        line=lineno,
                        context=_snippet(line, m.start(), m.end()),
                    )
                )
        for name, pattern in name_patterns:
            for m in pattern.finditer(line):
                findings.append(
                    Finding(
                        kind="denylisted_name",
                        match=m.group(0),
                        line=lineno,
                        context=_snippet(line, m.start(), m.end()),
                    )
                )
    return findings


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF using poppler's ``pdftotext``.

    Raises:
        RuntimeError: If pdftotext is unavailable or fails. Callers convert this
            into an "unscannable" finding so unscannable is never treated clean.
    """
    binary = shutil.which("pdftotext")
    if binary is None:
        raise RuntimeError("pdftotext (poppler) not available to extract PDF text")
    try:
        result = subprocess.run(
            [binary, "-q", str(path), "-"],
            capture_output=True,
            check=True,
        )
    except (subprocess.CalledProcessError, OSError) as exc:
        raise RuntimeError(f"pdftotext failed on {path}: {exc}") from exc
    return result.stdout.decode("utf-8", errors="replace")


def _extract_docx(path: Path) -> str:
    """Extract text (paragraphs + table cells) from a .docx via python-docx.

    Raises:
        RuntimeError: If python-docx is not importable or the file cannot be
            read. Callers convert this into an "unscannable" finding.
    """
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError(
            "python-docx not installed; cannot scan .docx files. "
            "Install with: pip install python-docx"
        ) from exc
    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001 - re-raised as clear RuntimeError below
        raise RuntimeError(f"could not open .docx {path}: {exc}") from exc
    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    """Extract all cell values from a .xlsx via openpyxl.

    Raises:
        RuntimeError: If openpyxl is not importable or the file cannot be read.
            Callers convert this into an "unscannable" finding.
    """
    try:
        import openpyxl  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError(
            "openpyxl not installed; cannot scan .xlsx files. "
            "Install with: pip install openpyxl"
        ) from exc
    try:
        workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001 - re-raised as clear RuntimeError below
        raise RuntimeError(f"could not open .xlsx {path}: {exc}") from exc
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                parts.extend(str(cell) for cell in row if cell is not None)
    finally:
        workbook.close()
    return "\n".join(parts)


def _read_text_file(path: Path) -> str:
    """Read a text file as UTF-8.

    Raises:
        RuntimeError: On decode failure, so binary content mislabeled as text is
            surfaced as unscannable rather than silently skipped.
    """
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise RuntimeError(f"could not decode {path} as UTF-8 text: {exc}") from exc


def _extract(path: Path) -> str:
    """Extract scannable text from a file based on its suffix.

    Raises:
        RuntimeError: When the appropriate extractor is unavailable or fails.
    """
    suffix = path.suffix.lower()
    if suffix in _TEXT_SUFFIXES:
        return _read_text_file(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".xlsx":
        return _extract_xlsx(path)
    # Unknown extension: attempt a UTF-8 read; failure becomes unscannable.
    return _read_text_file(path)


def scan_file(path: str | Path, denylist: list[str] | None = None) -> PIIReport:
    """Scan a single file for PII.

    Inputs:
        path: File to scan. Text formats (.md/.txt/.html/.csv) are read as
            UTF-8; .pdf is extracted with pdftotext; .docx via python-docx;
            .xlsx via openpyxl.
        denylist: Names to flag. If None, the bundled default is used.

    Output:
        A PIIReport. If text cannot be extracted (missing extractor, decode
        failure), a single "unscannable" Finding is recorded — unscannable is
        never reported as clean.

    Side effects:
        Reads from disk; may shell out to pdftotext for PDFs.
    """
    path = Path(path)
    try:
        text = _extract(path)
    except RuntimeError as exc:
        logger.warning("PII scan could not read %s: %s", path, exc)
        return PIIReport(
            path=str(path),
            findings=[Finding(kind="unscannable", match=str(exc), line=0, context=str(exc))],
        )
    return PIIReport(path=str(path), findings=scan_text(text, denylist=denylist))


def scan_paths(
    paths: Iterable[str | Path], denylist: list[str] | None = None
) -> list[PIIReport]:
    """Scan many files and return one PIIReport per path.

    Inputs:
        paths: An iterable of file paths.
        denylist: Names to flag. If None, the bundled default is used.

    Output:
        A list of PIIReport objects, one per input path, in input order.
    """
    return [scan_file(p, denylist=denylist) for p in paths]
